from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ============================================================
# 全局默认样式
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

pf = style.paragraph_format
pf.line_spacing = 1.25
pf.space_after = Pt(0)
pf.space_before = Pt(0)

# ============================================================
# 页面设置：A4 标准页边距
# ============================================================
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ============================================================
# 辅助函数
# ============================================================
def add_body_text(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    return p

def add_centered_text(text, bold=False, size=Pt(12), space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = size
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

# ============================================================
# 正文（精简后确保在一页内）
# ============================================================
add_centered_text('The Story of A4 Paper', bold=True, size=Pt(14), space_after=4)

add_body_text(
    'Every day, millions of people around the world use A4 paper for printing, copying, and drawing. '
    'Yet behind this common item lies a remarkable story of international standardization, scientific '
    'precision, and centuries of evolution.'
)

add_body_text(
    'Before standardization, paper sizes varied wildly between countries and even between paper mills. '
    'In 1922, Germany solved this chaos by publishing DIN 476, which defined paper sizes based on an '
    'elegant principle: an aspect ratio of √2:1. When a sheet with this ratio is cut in half, each half '
    'retains the exact same proportions — a property called self-similarity. Interestingly, this idea '
    'was first proposed by German scientist Georg Christoph Lichtenberg as early as 1786.'
)

add_body_text(
    'DIN 476 became the foundation for the international standard ISO 216, published in 1975. It defines '
    'three series of paper sizes: A, B, and C. In the A series, A0 measures 841 mm × 1189 mm (exactly one '
    'square meter), and each subsequent size is half the area of the previous one. A4 paper (210 mm × 297 mm) '
    'has become the most practical size for everyday use — large enough for text and graphics, yet easy to '
    'handle and file.'
)

add_body_text(
    'This standardization had a profound global impact. It eliminated confusion and waste from incompatible '
    'sizes, making international correspondence seamless. Printers, photocopiers, and filing cabinets can now '
    'be manufactured to a single standard, creating economies of scale and lower costs. Today, ISO 216 is '
    'adopted by nearly all countries except the United States and Canada, which still use letter-size paper.'
)

add_body_text(
    'The production of A4 paper also raises environmental concerns, as the paper industry consumes millions '
    'of trees and vast amounts of water and energy. In response, many manufacturers offer recycled A4 paper, '
    'and certification schemes like FSC promote sustainable forestry. Although the digital age has reduced '
    'paper use, demand remains strong in education and legal sectors.'
)

add_body_text(
    'The story of A4 paper shows that even the most ordinary objects have extraordinary histories. From '
    'Lichtenberg\'s insight to the DIN standard and global ISO 216, A4 paper represents a remarkable '
    'achievement of human ingenuity and international cooperation.'
)

# ============================================================
# 保存文档
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'The Story of A4 Paper.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
